"""
CERBERUS — Python AI Backend
Port: 8000 (Bound to 127.0.0.1 — localhost only)

Role: Pure AI analysis service (NO x402 here — that lives in the Node x402-server).
The x402 Node gateway server (port 4021) verifies payment, then proxies to this backend.

Security:
  - Binds to 127.0.0.1 ONLY — not reachable externally.
  - Direct calls to /analyze-contract blocked unless X-Internal-Secret matches .env secret.
  - INTERNAL_SECRET has NO hardcoded default — server refuses to start if unset.
  - CORS locked to ALLOWED_ORIGIN env var (default: the gateway origin only).
  - All error responses stripped of internal stack traces.
"""

import json
import re
import base64
import io
import zipfile
import logging
import os
import time
from collections import defaultdict
from typing import List, Optional
from fastapi import FastAPI, HTTPException, Request, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
import httpx

# Optional document parsers — server-side text extraction fallback for uploads
try:
    import pypdf  # PDF text extraction
except ImportError:
    pypdf = None
try:
    import docx  # python-docx (.docx)
except ImportError:
    docx = None

# Load backend/.env — INTERNAL_SECRET must come from .env (NO hardcoded default).
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ─── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)-8s %(message)s",
    datefmt="%Y-%m-%dT%H:%M:%S",
)
logger = logging.getLogger(__name__)

# ─── Config ───────────────────────────────────────────────────────────────────
OLLAMA_BASE_URL  = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL     = os.getenv("OLLAMA_MODEL", "mistral:latest")
OLLAMA_TIMEOUT   = int(os.getenv("OLLAMA_TIMEOUT", "180"))

# ALLOWED_ORIGIN — in production, lock this to the gateway's origin (port 4021).
# Defaults to the local gateway to prevent the backend from accepting browser requests directly.
ALLOWED_ORIGIN   = os.getenv("ALLOWED_ORIGIN", "http://localhost:4021")

# INTERNAL_SECRET — NO hardcoded fallback.
INTERNAL_SECRET = os.getenv("INTERNAL_SECRET")
if not INTERNAL_SECRET:
    import sys
    print(
        "FATAL: INTERNAL_SECRET env var is not set. "
        "Set it in backend/.env. Refusing to start.",
        flush=True,
    )
    sys.exit(1)

# ─── Metrics counters (in-memory) ─────────────────────────────────────────────
_start_time    = time.time()
_request_count = 0
_error_count   = 0
_analysis_count = 0

# ─── FastAPI App ──────────────────────────────────────────────────────────────
app = FastAPI(
    title="CERBERUS — AI Backend",
    version="2.0.0",
    docs_url=None,   # Disable Swagger UI in production (no public API docs)
    redoc_url=None,
)

# CORS — only the gateway should talk to this backend.
# The gateway runs on localhost:4021 and proxies all frontend requests.
app.add_middleware(
    CORSMiddleware,
    allow_origins=[ALLOWED_ORIGIN],
    allow_credentials=False,
    allow_methods=["POST", "GET"],
    allow_headers=["Content-Type", "X-Internal-Secret"],
)

# ─── Models ───────────────────────────────────────────────────────────────────
class AnalyzeRequest(BaseModel):
    contract_text: str


class ExtractRequest(BaseModel):
    file_name: str = "contract.txt"
    file_b64: str = ""  # base64-encoded file bytes


# ─── Limits ───────────────────────────────────────────────────────────────────
MAX_FILE_BYTES = 20 * 1024 * 1024  # 20 MB upload cap
MAX_TEXT_CHARS  = 50_000             # matches /analyze-contract limit


# ─── Internal Secret Checker ──────────────────────────────────────────────────
def _require_internal_secret(x_internal_secret: Optional[str]) -> None:
    """Raise 403 if the gateway secret header is missing or wrong."""
    if x_internal_secret != INTERNAL_SECRET:
        # Use timing-safe comparison to prevent timing attacks
        import hmac
        valid = x_internal_secret is not None and hmac.compare_digest(
            x_internal_secret.encode(), INTERNAL_SECRET.encode()
        )
        if not valid:
            logger.warning("Direct access attempt blocked (missing/invalid X-Internal-Secret)")
            raise HTTPException(
                status_code=403,
                detail="Forbidden: All requests must be routed through the CERBERUS x402 payment gateway.",
            )


# ─── Document Text Extraction ─────────────────────────────────────────────────
def extract_text_from_bytes(filename: str, data: bytes) -> str:
    """Best-effort text extraction for common contract document formats."""
    name = (filename or "").lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""

    # PDF
    if ext == "pdf":
        if pypdf is None:
            raise ValueError("PDF support not installed on server (pip install pypdf)")
        try:
            reader = pypdf.PdfReader(io.BytesIO(data))
            pages = [page.extract_text() or "" for page in reader.pages]
        except Exception:
            raise ValueError(
                "Could not read PDF — the file may be corrupt, password-protected, "
                "or a scanned image without a text layer."
            )
        return "\n\n".join(pages)

    # Word (.docx)
    if ext == "docx":
        if docx is None:
            raise ValueError("DOCX support not installed on server (pip install python-docx)")
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    # Word (.doc — legacy binary)
    if ext == "doc":
        raw = data.decode("cp1252", errors="ignore")
        raw = re.sub(r"[^\x20-\x7E\r\n\t]", " ", raw)
        raw = re.sub(r"\s{3,}", "\n", raw)
        if len(raw.strip()) < 20:
            raise ValueError("Could not extract text from legacy .doc (save as .docx or .pdf and retry)")
        return raw

    # OpenDocument Text (.odt / .fodt)
    if ext in ("odt", "fodt"):
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                xml = z.read("content.xml").decode("utf-8", errors="ignore")
            xml = re.sub(r"<text:p[^>]*>", "\n", xml)
            xml = re.sub(r"<[^>]+>", "", xml)
            return re.sub(r"\n{3,}", "\n\n", xml)
        except Exception:
            raise ValueError("Could not parse .odt document")

    # RTF
    if ext == "rtf":
        if not data.startswith(b"{\\rtf"):
            raise ValueError("Not a valid RTF file")
        raw = data.decode("cp1252", errors="ignore")
        raw = re.sub(r"\\par[d]?", "\n", raw)
        raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", raw)
        raw = raw.replace("{", "").replace("}", "")
        return re.sub(r"\n{3,}", "\n\n", raw)

    # HTML / XML
    if ext in ("html", "htm", "xml"):
        import html as html_mod
        text = html_mod.unescape(data.decode("utf-8", errors="ignore"))
        text = re.sub(r"<script[^>]*>.*?</script>", " ", text, flags=re.S | re.I)
        text = re.sub(r"<style[^>]*>.*?</style>", " ", text, flags=re.S | re.I)
        text = re.sub(r"<[^>]+>", " ", text)
        return re.sub(r"[ \t]{2,}", " ", text)

    # Plain text family
    if ext in ("txt", "md", "markdown", "csv", "tsv", "json", "log", "text"):
        return data.decode("utf-8", errors="replace")

    # Unknown — sniff for readable UTF-8
    try:
        return data.decode("utf-8", errors="strict")
    except UnicodeDecodeError:
        pass
    raw = data.decode("cp1252", errors="ignore")
    raw = re.sub(r"[^\x20-\x7E\r\n\t]", "", raw)
    if len(raw.strip()) < 20:
        raise ValueError(
            f"Could not extract text from '{filename}'. "
            "Supported formats: PDF, Word (.doc/.docx), ODT, RTF, HTML, and plain text."
        )
    return raw


# ─── Sample Contract ──────────────────────────────────────────────────────────
SAMPLE_CONTRACT = """SERVICE AGREEMENT

1. INDEMNIFICATION
The Client shall indemnify and hold harmless the Service Provider, its officers, directors,
employees, and agents from and against any and all claims, damages, losses, costs, and expenses
(including reasonable attorneys' fees) arising out of or relating to this Agreement, regardless
of whether such claims arise from the negligence or willful misconduct of the Service Provider.

2. LIMITATION OF LIABILITY
In no event shall the Service Provider be liable for any indirect, incidental, special,
consequential, or punitive damages, or any loss of profits or revenues. The Service Provider's
total liability shall not exceed $100, regardless of the nature of the claim.

3. INTELLECTUAL PROPERTY ASSIGNMENT
The Client agrees that all work product, inventions, discoveries, and creations developed
by the Client, even during personal time and using personal resources, shall be the sole
property of the Service Provider. The Client hereby assigns all intellectual property rights
to the Service Provider in perpetuity.

4. TERMINATION WITHOUT CAUSE
The Service Provider may terminate this Agreement at any time, for any reason or no reason,
with zero notice, and without any obligation to pay severance or any outstanding compensation.

5. GOVERNING LAW
This Agreement shall be governed by the laws of a jurisdiction chosen solely at the Service
Provider's discretion, which may be changed at any time without notice to the Client.

6. NON-COMPETE CLAUSE
For a period of 5 years after termination, the Client agrees not to work in any capacity
in any industry that competes with the Service Provider's current or future business interests
anywhere in the world."""

SYSTEM_PROMPT = """You are a legal risk analysis AI. Analyze contract text and identify risky clauses.

OUTPUT RULES (STRICT):
- Output ONLY a valid JSON array. NO markdown fences, NO explanations, NO preamble.
- Start your response with [ and end with ]
- Each object must have exactly these keys: clause, risk_level, reason, suggested_rewrite
- risk_level must be exactly: "high", "medium", or "low"
- clause: the specific problematic text quoted from the contract (max 150 chars)
- reason: why this is risky (1-2 sentences)
- suggested_rewrite: a safer alternative (1-3 sentences)

Example of correct output format:
[{"clause":"unlimited indemnification clause","risk_level":"high","reason":"Exposes client to unlimited liability.","suggested_rewrite":"Limit indemnification to direct damages caused by client negligence."}]"""


def extract_json_array(raw: str) -> list:
    """Robustly extract a JSON array even if model wraps in markdown/prose."""
    raw = raw.strip()
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, list):
            return parsed
    except json.JSONDecodeError:
        pass

    clean = re.sub(r"```(?:json)?", "", raw).strip()
    try:
        parsed = json.loads(clean)
        if isinstance(parsed, list):
            return parsed
    except json.JSONDecodeError:
        pass

    start = raw.find("[")
    end   = raw.rfind("]")
    if start != -1 and end != -1 and end > start:
        try:
            parsed = json.loads(raw[start : end + 1])
            if isinstance(parsed, list):
                return parsed
        except json.JSONDecodeError:
            pass

    raise ValueError(f"Could not parse JSON array from model output: {raw[:300]}")


async def analyze_with_ollama(contract_text: str) -> List[dict]:
    logger.info(f"Sending to Ollama ({OLLAMA_MODEL}), length={len(contract_text)}")
    payload = {
        "model": OLLAMA_MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": f"Analyze this contract and return JSON array:\n\n{contract_text}"},
        ],
        "stream": False,
        "options": {
            "temperature": 0.1,
            "num_predict": 3000,
        },
    }

    async with httpx.AsyncClient(timeout=OLLAMA_TIMEOUT) as client:
        resp = await client.post(f"{OLLAMA_BASE_URL}/api/chat", json=payload)
        resp.raise_for_status()

    data        = resp.json()
    raw_content = data["message"]["content"]
    clauses     = extract_json_array(raw_content)

    normalised = []
    for item in clauses:
        risk = str(item.get("risk_level", "medium")).lower().strip()
        if risk not in ("high", "medium", "low"):
            risk = "medium"
        normalised.append({
            "clause":            str(item.get("clause", ""))[:500],
            "risk_level":        risk,
            "reason":            str(item.get("reason", "")),
            "suggested_rewrite": str(item.get("suggested_rewrite", "")),
        })

    logger.info(f"Extracted {len(normalised)} risky clauses")
    return normalised


# ─── Request counter middleware ────────────────────────────────────────────────
@app.middleware("http")
async def count_requests(request: Request, call_next):
    global _request_count, _error_count
    _request_count += 1
    response = await call_next(request)
    if response.status_code >= 500:
        _error_count += 1
    return response


# ─── Public Routes ─────────────────────────────────────────────────────────────
@app.get("/health")
async def health():
    return {
        "status": "ok",
        "service": "CERBERUS-AI",
        "version": "2.0.0",
        "model": OLLAMA_MODEL,
    }


@app.get("/metrics")
async def metrics():
    """Operational metrics for monitoring dashboards."""
    return {
        "uptime_seconds": round(time.time() - _start_time),
        "requests_total": _request_count,
        "errors_total": _error_count,
        "analyses_total": _analysis_count,
        "model": OLLAMA_MODEL,
    }


@app.get("/sample-contract")
async def get_sample():
    return {"contract_text": SAMPLE_CONTRACT.strip()}


# ─── Protected AI Endpoints ────────────────────────────────────────────────────

@app.post("/analyze-contract")
async def analyze_contract(
    req: AnalyzeRequest,
    x_internal_secret: Optional[str] = Header(None, alias="X-Internal-Secret"),
):
    """
    AI analysis endpoint — ONLY callable via CERBERUS x402 gateway (port 4021).
    Direct external calls without valid X-Internal-Secret header return 403.
    """
    global _analysis_count
    _require_internal_secret(x_internal_secret)

    text = req.contract_text.strip()
    if not text:
        raise HTTPException(400, "contract_text is required")
    if len(text) > 50_000:
        raise HTTPException(400, f"Contract too long ({len(text):,} chars). Maximum is 50,000 characters.")

    logger.info("=== CERBERUS /analyze-contract called via x402 proxy ===")
    t0 = time.time()

    try:
        results = await analyze_with_ollama(text)
    except httpx.ConnectError:
        logger.error("Ollama not reachable")
        raise HTTPException(503, "AI engine is not available. Please ensure Ollama is running.")
    except httpx.HTTPStatusError as e:
        logger.error(f"Ollama HTTP error: {e.response.status_code}")
        raise HTTPException(503, "AI engine returned an error. Please try again.")
    except httpx.TimeoutException:
        logger.error("Ollama timeout")
        raise HTTPException(504, "AI analysis timed out. Try with a shorter contract or retry.")
    except ValueError as e:
        logger.error(f"JSON parse error: {e}")
        raise HTTPException(500, "AI returned an unreadable response. Please retry.")

    _analysis_count += 1
    elapsed = round(time.time() - t0, 2)
    logger.info(f"=== Analysis complete: {len(results)} clauses in {elapsed}s ===")
    return results


@app.post("/extract-text")
async def extract_text(
    req: ExtractRequest,
    x_internal_secret: Optional[str] = Header(None, alias="X-Internal-Secret"),
):
    """
    Server-side text extraction for uploaded contract documents.
    ONLY callable via the CERBERUS x402 gateway (port 4021).

    Request:  { file_name: "contract.pdf", file_b64: "<base64>" }
    Response: { text, chars, format }
    """
    _require_internal_secret(x_internal_secret)

    try:
        data = base64.b64decode(req.file_b64, validate=True)
    except Exception:
        raise HTTPException(400, "file_b64 must be valid base64 data")

    if not data:
        raise HTTPException(400, "File is empty")
    if len(data) > MAX_FILE_BYTES:
        raise HTTPException(413, f"File too large (max {MAX_FILE_BYTES // (1024 * 1024)} MB)")

    # Sanitize filename — prevent path traversal
    safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", req.file_name)
    logger.info(f"=== /extract-text: {safe_name} ({len(data):,} bytes) ===")

    try:
        text = extract_text_from_bytes(safe_name, data)
    except ValueError as e:
        logger.warning(f"Extraction failed for {safe_name}: {e}")
        raise HTTPException(422, str(e))
    except Exception:
        logger.exception(f"Unexpected extraction error for {safe_name}")
        raise HTTPException(500, "Could not extract text from the uploaded document.")

    if len(text) > MAX_TEXT_CHARS:
        marker = "\n\n[...TRUNCATED — analysis limited to 50,000 chars]"
        text = text[: MAX_TEXT_CHARS - len(marker)] + marker

    return {
        "text":   text,
        "chars":  len(text),
        "format": safe_name.rsplit(".", 1)[-1].lower() if "." in safe_name else "text",
    }


# ─── Static Frontend Serving ───────────────────────────────────────────────────
FRONTEND_DIR = os.path.join(os.path.dirname(__file__), "..", "frontend")
if os.path.isdir(FRONTEND_DIR):
    app.mount("/static", StaticFiles(directory=FRONTEND_DIR), name="static")

    @app.get("/")
    async def serve_frontend():
        return FileResponse(os.path.join(FRONTEND_DIR, "index.html"))


if __name__ == "__main__":
    # If run directly (`python main.py`), ALWAYS bind to localhost only.
    import uvicorn
    uvicorn.run(
        app,
        host="127.0.0.1",
        port=8000,
        log_level="info",
        access_log=True,
    )
